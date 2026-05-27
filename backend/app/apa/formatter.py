"""Aplica un perfil APA a un documento Word usando python-docx."""

from __future__ import annotations

import re
from typing import Optional

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from .profiles import ApaProfile, HeadingStyle

# Encabezados (en espanol e ingles) que marcan el inicio de la lista de
# referencias. A partir de aqui se aplica sangria francesa.
_REFERENCE_TITLES = {
    "referencias",
    "referencia",
    "bibliografia",
    "bibliografía",
    "lista de referencias",
    "references",
    "reference list",
    "works cited",
}

_ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# Estilos de titulo de Word, en ingles ("Heading 1") y espanol ("Titulo 1").
_HEADING_RE = re.compile(r"^(?:heading|t[ií]tulo|titulo)\s*(\d)$", re.IGNORECASE)


def apply_profile(doc: DocxDocument, profile: ApaProfile) -> None:
    """Aplica las reglas del perfil al documento, modificandolo en sitio."""
    _apply_margins(doc, profile)
    _apply_base_style(doc, profile)
    _add_page_numbers(doc, profile)

    headings = {h.level: h for h in profile.headings}
    in_references = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        level = _heading_level(paragraph)

        if text.lower() in _REFERENCE_TITLES:
            in_references = True
            _format_references_title(paragraph, profile)
            continue

        if level is not None and level in headings:
            _format_heading(paragraph, headings[level], profile)
            continue

        if in_references and text:
            _format_reference_entry(paragraph, profile)
            continue

        _format_body(paragraph, profile)


def _apply_margins(doc: DocxDocument, profile: ApaProfile) -> None:
    margin = Inches(profile.margin_in)
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def _apply_base_style(doc: DocxDocument, profile: ApaProfile) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = profile.font_name
    normal.font.size = Pt(profile.font_size_pt)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _add_page_numbers(doc: DocxDocument, profile: ApaProfile) -> None:
    """Numero de pagina arriba a la derecha en todas las paginas (APA)."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        # Limpia el contenido previo del encabezado y deja solo el numero.
        for para in list(header.paragraphs):
            para._element.getparent().remove(para._element)
        para = header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run()
        _set_run_font(run, profile)
        _append_page_field(run)


def _append_page_field(run) -> None:
    """Inserta un campo de Word { PAGE } en el run (lo evalua Word/LibreOffice)."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    # Valor en cache para que el numero se vea aunque no se actualicen campos.
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._element.append(begin)
    run._element.append(instr)
    run._element.append(cached)
    run._element.append(end)


def _heading_level(paragraph: Paragraph) -> Optional[int]:
    style = paragraph.style
    name = (style.name or "") if style else ""
    match = _HEADING_RE.match(name.strip())
    return int(match.group(1)) if match else None


def _set_run_font(run, profile: ApaProfile) -> None:
    run.font.name = profile.font_name
    run.font.size = Pt(profile.font_size_pt)
    # Fija la fuente en todos los scripts (ascii, hAnsi, cs, eastAsia) para que
    # no quede texto con la fuente original por herencia.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), profile.font_name)


def _format_body(paragraph: Paragraph, profile: ApaProfile) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.alignment = _ALIGNMENTS.get(profile.body_alignment, WD_ALIGN_PARAGRAPH.LEFT)
    pf.first_line_indent = Inches(profile.first_line_indent_in)
    pf.left_indent = Inches(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    # El formato directo de los runs pisa al estilo Normal, asi que se reaplica.
    for run in paragraph.runs:
        _set_run_font(run, profile)


def _format_reference_entry(paragraph: Paragraph, profile: ApaProfile) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Sangria francesa: el bloque sangra y la primera linea vuelve al margen.
    pf.left_indent = Inches(profile.hanging_indent_in)
    pf.first_line_indent = Inches(-profile.hanging_indent_in)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for run in paragraph.runs:
        _set_run_font(run, profile)


def _format_references_title(paragraph: Paragraph, profile: ApaProfile) -> None:
    pf = paragraph.paragraph_format
    # La lista de referencias empieza en una pagina nueva.
    pf.page_break_before = True
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.alignment = _ALIGNMENTS.get(
        profile.references_title_alignment, WD_ALIGN_PARAGRAPH.CENTER
    )
    pf.first_line_indent = Inches(0)
    pf.left_indent = Inches(0)
    for run in paragraph.runs:
        _set_run_font(run, profile)
        run.bold = profile.references_title_bold


def _format_heading(
    paragraph: Paragraph, heading: HeadingStyle, profile: ApaProfile
) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.alignment = _ALIGNMENTS.get(heading.alignment, WD_ALIGN_PARAGRAPH.LEFT)
    pf.first_line_indent = Inches(profile.first_line_indent_in if heading.inline else 0)
    pf.left_indent = Inches(0)
    for run in paragraph.runs:
        _set_run_font(run, profile)
        run.bold = heading.bold
        run.italic = heading.italic
